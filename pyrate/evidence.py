"""
Evidence generation utilities for PyRate Framework.

Automatically generates evidence files for test executions:
- API tests: TXT files with request/response details
- UI tests: DOCX files with screenshots and step descriptions
"""

import os
import json
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, RGBColor
from io import BytesIO


class EvidenceGenerator:
    """
    Generates test evidence files for API and UI automation.
    
    API evidence is saved as TXT files with request/response information.
    UI evidence is saved as DOCX files with screenshots and step details.
    
    Attributes:
        output_folder: Directory where evidence files are saved
    """
    
    def __init__(self, output_folder: str = "evidence"):
        """
        Initialize evidence generator.
        
        Args:
            output_folder: Directory for evidence files (created if doesn't exist)
        """
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)

    def generate_api_evidence(
        self, 
        scenario_name: str, 
        method: str, 
        response_data: Any, 
        iteration: int = 0
    ) -> str:
        """
        Generate TXT evidence file for API test execution.
        
        Args:
            scenario_name: Name of the test scenario
            method: HTTP method used (GET, POST, etc.)
            response_data: Response data (usually dict or string)
            iteration: Data iteration number (for data-driven tests)
            
        Returns:
            Path to the generated evidence file
            
        Example:
            >>> gen = EvidenceGenerator()
            >>> path = gen.generate_api_evidence(
            ...     "Login Test", 
            ...     "POST", 
            ...     {"token": "abc123"},
            ...     iteration=1
            ... )
        """
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        safe_name = "".join([c if c.isalnum() else "_" for c in scenario_name])
        filename = f"{self.output_folder}/API_{safe_name}_Iter{iteration + 1}.txt"

        content = f"""
========================================
REPORTE DE EJECUCIÓN API
========================================
Fecha/Hora: {timestamp}
Escenario:  {scenario_name}
Método:     {method}
Iteración:  {iteration + 1}
========================================

RESPONSE (JSON):
----------------
{json.dumps(response_data, indent=4, ensure_ascii=False)}
        """

        with open(filename, "w", encoding="utf-8") as f:
            f.write(content.strip())

        return filename

    def generate_ui_evidence(
        self, 
        scenario_name: str, 
        steps_log: List[Dict[str, Any]], 
        iteration: int = 0
    ) -> str:
        """
        Generate DOCX evidence file for UI test execution with screenshots.
        
        Args:
            scenario_name: Name of the test scenario
            steps_log: List of execution steps with screenshots
            iteration: Data iteration number (for data-driven tests)
            
        Returns:
            Path to the generated DOCX file
            
        Example:
            >>> steps = [
            ...     {
            ...         "name": "Click login button",
            ...         "status": "PASS",
            ...         "screenshot_bytes": b"...",
            ...         "error": None
            ...     }
            ... ]
            >>> path = gen.generate_ui_evidence("Login Flow", steps)
        """
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        safe_name = "".join([c if c.isalnum() else "_" for c in scenario_name])
        filename = f"{self.output_folder}/UI_{safe_name}_Iter{iteration + 1}.docx"

        doc = Document()

        # Título
        doc.add_heading(f'Evidencia de Prueba: {scenario_name}', 0)

        # Metadatos
        p = doc.add_paragraph()
        p.add_run('Fecha de Ejecución: ').bold = True
        p.add_run(timestamp + '\n')
        p.add_run('Iteración de Datos: ').bold = True
        p.add_run(str(iteration + 1))

        doc.add_heading('Detalle Paso a Paso', level=1)

        # Iterar pasos
        for step in steps_log:
            # Texto del paso
            p_step = doc.add_paragraph()
            run_step = p_step.add_run(f"Paso: {step['name']}")
            run_step.bold = True

            # Estado y Colores
            if step['status'] == 'FAIL':
                run_status = p_step.add_run(f" [FALLIDO]")
                run_status.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
                if step.get('error'):
                    doc.add_paragraph(f"Error: {step['error']}")
            else:
                run_status = p_step.add_run(f" [OK]")
                run_status.font.color.rgb = RGBColor(0, 128, 0)  # Verde

            # Respuesta API (si hubo)
            if step.get('response_data'):
                doc.add_paragraph("Respuesta API:", style='Caption')
                doc.add_paragraph(str(step['response_data'])[:1000])

            # Screenshot (si hubo)
            if step.get('screenshot_bytes'):
                doc.add_paragraph("Evidencia Visual:", style='Caption')
                try:
                    image_stream = BytesIO(step['screenshot_bytes'])
                    doc.add_picture(image_stream, width=Inches(5.0))
                except Exception as e:
                    doc.add_paragraph(f"[No se pudo adjuntar imagen: {str(e)}]")

            doc.add_paragraph("-" * 50)

        doc.save(filename)
        return filename